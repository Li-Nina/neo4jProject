import gzip
import logging
import csv
import tarfile
import zipfile
import rarfile

from elasticsearch import Elasticsearch
from elasticsearch_dsl import Search
from elasticsearch_dsl import connections
from config import CLUSTER_NODES
import pymysql
import openpyxl
import os
import docx
import time



logger = logging.getLogger(__name__)
logger.warning("getting connection to es ...")
es = Elasticsearch([CLUSTER_NODES])
my_index = 'con_pro_2'
my_doc_type = '_doc'


# connections.create_connection(hosts=[CLUSTER_NODES], timeout=20)
# client = Elasticsearch()

#
# def searchExpert():
#     s = Search(using=client, index="my-index") \
#         .filter("term", category="search") \
#         .query("match", title="python") \
#         .exclude("match", description="beta")
#
#     s.aggs.bucket('per_tag', 'terms', field='tags') \
#         .metric('max_lines', 'max', field='lines')
#
#     response = s.execute()
#
#     for hit in response:
#         print(hit.meta.score, hit.title)
#
#     for tag in response.aggregations.per_tag.buckets:
#         print(tag.key, tag.max_lines.value)
# -*- coding:UTF-8 -*-
# 创建索引
def create_indexs_es():
    _index_mappings = {
        "mappings": {
            "dynamic": "true",  # 动态的
            "properties": {
                "ExpertID": {
                    "type": "text",
                    "index": False
                },
                "fieldID": {
                    "type": "text",
                    "index": False
                },
                "name": {
                    "type": "text"
                },
                "college": {
                    "type": "text",
                    "analyzer": "ik_max_word"
                },
                "education": {
                    "type": "text",
                    "analyzer": "ik_max_word"
                },
                "degree": {
                    "type": "text",
                    "analyzer": "ik_max_word"
                },
                "job": {
                    "type": "text",
                    "analyzer": "ik_max_word"
                },
                "subject": {
                    "type": "text",
                    "analyzer": "ik_max_word"
                },
                "phone": {
                    "type": "text"
                },

                "address": {
                    "type": "text",
                    "analyzer": "ik_smart"  # 粗粒度分词
                },
                "research": {
                    "type": "text",
                    "analyzer": "ik_max_word"
                },

                "department": {
                    "type": "text",
                    "analyzer": "ik_max_word"
                },
                "title": {
                    "type": "text",
                    "analyzer": "ik_max_word"
                },
                "email": {
                    "type": "text"
                },
                "field": {
                    "type": "text",
                    "analyzer": "ik_max_word"
                },
                "project": {
                    "type": "text",
                    "analyzer": "ik_smart"  # 粗粒度分词
                }
            }
        }
    }
    if es.indices.exists(index=my_index) is not True:
        res = es.indices.create(index=my_index, body=_index_mappings)
        print("无索引，创建它:\n", res)
    else:
        print("索引已经存在，跳过创建！")


# 删除索引
def delete_index():
    if es.indices.exists(index=my_index) is True:
        res = es.indices.delete(index=my_index)


# 从CSV文件中读取数据，并存储到es中
def index_data_from_csv(csvfile):
    """
    :param csvfile: csv文件，包括完整路径
    """
    # 取CSV数据
    with open(csvfile, 'r', encoding='utf-8', newline='') as f:
        list_data = csv.reader(f)
        # 循环成字典形式，插入数据
        index = 0

        for item in list_data:
            doc = {}
            if index > 0:  # 第一行是标题
                doc['Expert_ID'] = item[0]
                # doc['fieldID'] = item[1]
                doc['name'] = item[2]
                doc['college'] = item[3]
                doc['education'] = item[4]
                doc['degree'] = item[5]
                doc['job'] = item[6]
                # doc['职称'] = item[7]
                doc['subject'] = item[8]
                doc['phone'] = item[9]
                doc['address'] = item[10]
                doc['research'] = item[11]
                doc['department'] = item[12]
                doc['title'] = item[13]
                doc['email'] = item[14]
                doc['field'] = item[15]
                doc['project'] = item[16]
                # print(doc)
                res = es.index(index=my_index, doc_type=my_doc_type, body=doc,id=doc['Expert_ID'])
                # print(res)
            index += 1
            # print(index)


# 根据请求体查询
def get_data_by_body(input):
    # dsl = {
    #     "query": {
    #         "match": {
    #             "领域": "区块链"
    #         }
    #     }
    # }
    dsl = {
        "query": {
            "multi_match": {
                "query": input,
                "type": "best_fields",
                # "type": "most_fields",
                # "fields" : ["学历", "领域",'单位','职务','职称','学科','研究方向','身份','领域','项目']
                "fields": ['college', 'education', 'degree', 'job', 'field', 'subject', 'research', 'department',
                           'title', 'project'],
                "tie_breaker": 0.3

            }
        }

    }
    res = es.search(index=my_index, body=dsl)
    return res
#
def get_data_from_mysql():
    """
    从MySQL数据库中读取数据
    :return: 元组
    """
    try:
        # 2.建立与数据库连接 TODO: 用户名root，密码x5
        conn = pymysql.connect(host="localhost", user="root", password="x5", database="expertserverbig",
                               port=3306,
                               charset="UTF8")

        # 3.获取游标
        cursor = conn.cursor()
    except ConnectionError as e:
        print("数据库连接异常 或 获取游标异常！", e)
        return []

    try:
        # TODO: Java组的数据没有expert_id这些数据，所以加上了expert_id is not null
        select_sql = "select expert_id, fieldId, name, college, education, degree, job, professional, subject, " \
                     "phonenumber, address, research, department, title, email, field, project, is_download " \
                     "from texpert_info " \
                     "where expert_id is not null and is_delete=FALSE order by publish_time desc, update_time desc;"
        rows = cursor.execute(select_sql)
        if rows == 0:
            print('数据库没有数据')
            return
        # 移动指针，0代表移动几个位置，absolute表示回到头部，即：从头部移动0个位置cursor.scroll(0, 'absolute')
        # cursor.scroll(0, 'absolute')
        # 取出所有数据，返回的格式是：((), (),...) 元组套元组
        data_all = cursor.fetchall()

        # 将数据拼接成[{key:value}]的形式
        mysql_list = []
        for item in data_all:
            doc = {}
            doc['Expert_ID'] = item[0]  # 专家id  +
            doc['fieldID'] = item[1]  # +
            doc['name'] = item[2]  # name
            doc['college'] = item[3]  # 单位  +
            doc['education'] = item[4]  # 学历  +
            doc['degree'] = item[5]  # 学位  +
            doc['job'] = item[6]  # 职务  +
            # doc['职称'] = item[7]  # professional
            doc['subject'] = item[8]  # 学科  +
            doc['phone'] = item[9]  # 工作电话  phonenumber
            doc['address'] = item[10]  # 地址  address
            doc['research'] = item[11]  # 研究方向  +
            doc['department'] = item[12]  # 所属部门  +
            doc['title'] = item[13]  # 身份  +
            doc['email'] = item[14]  # 邮箱 email
            doc['field'] = item[15]  # 领域  +
            doc['project'] = item[16]  # 项目  +
            if not item[17]:
                doc['is_download'] = True  # 是否可以下载
            else:
                doc['is_download'] = False
            mysql_list.append(doc)
        print(mysql_list)
        return mysql_list

    except pymysql.MySQLError as e:
        print('数据库操作异常，开始回滚', e)
        conn.rollback()
        return []
    finally:
        # 5：关闭游标、关闭连接
        cursor.close()
        conn.close()

# 导入数据到es中
def import_index_data(index_data):
    for data in index_data:
        es.index(index=my_index, doc_type=my_doc_type, body=data, id=data['Expert_ID'])


# 解析Excel文件
def parse_excel(excel_path):
    """
    excel_path: excel文件路径
    return: 解析后的字典
    """
    my_book = openpyxl.load_workbook(excel_path)
    my_sheet = my_book['Sheet1']

    # 开始解析数据
    doc = {}
    doc['Expert_ID'] = str(my_sheet.cell(2, 6).value)
    doc['fieldID'] = my_sheet.cell(2, 8).value
    doc['name'] = my_sheet.cell(3, 2).value
    doc['college'] = my_sheet.cell(7, 2).value
    doc['education'] = my_sheet.cell(4, 2).value
    doc['degree'] = my_sheet.cell(4, 4).value
    doc['job'] = my_sheet.cell(4, 6).value
    doc['professional'] = my_sheet.cell(4, 8).value
    doc['subject'] = my_sheet.cell(5, 2).value
    doc['phone'] = my_sheet.cell(6, 2).value
    doc['address'] = my_sheet.cell(8, 2).value
    doc['research'] = my_sheet.cell(9, 2).value
    doc['department'] = my_sheet.cell(5, 6).value
    doc['title'] = my_sheet.cell(3, 8).value
    doc['email'] = my_sheet.cell(6, 6).value
    doc['field'] = my_sheet.cell(5, 8).value
    doc['project'] = my_sheet.cell(11, 2).value
    return doc

# 解析word文件
word_path = ''
def parse_word(word_path):
    """
    excel_path: excel文件路径
    return: 解析后的字典
    """
    file = docx.Document(word_path)
    t = file.tables[0]
    # 开始解析数据
    doc = {}
    expert_id = file.paragraphs[1].text
    doc['Expert_ID'] = expert_id.split(':')[1]
    field_id = file.paragraphs[2].text
    doc['fieldID'] =  field_id.split(':')[1]
    doc['name'] = t.cell(0, 1).text
    doc['title'] = t.cell(0, 7).text
    doc['education'] = t.cell(1, 1).text
    doc['degree'] = t.cell(1, 3).text
    doc['job'] = t.cell(1, 5).text
    doc['professional'] = t.cell(1, 7).text
    doc['subject'] = t.cell(2, 1).text
    doc['college'] = t.cell(2, 3).text
    doc['email'] = t.cell(2, 5).text
    doc['field'] = t.cell(3, 1).text
    doc['department'] = t.cell(3, 3).text
    doc['phone'] = t.cell(3, 5).text
    doc['address'] = t.cell(4, 1).text
    doc['research'] = t.cell(5, 1).text
    doc['project'] = t.cell(6, 1).text
    return doc

# 保存到数据库
# 保存到数据库
def save_to_mysql(data):
    try:
        # 2.建立与数据库连接 TODO: 用户名root，密码x5
        conn = pymysql.connect(host="localhost",
                               user="root",
                               password="x5",
                               database="expertserverbig",
                               port=3306,
                               charset="UTF8")

        # 3.获取游标
        cursor = conn.cursor()
    except ConnectionError as e:
        print("数据库连接异常 或 获取游标异常！", e)
        return "数据保存失败"

    try:
        if len(data.keys()) == 1 and "Expert_ID" in data.keys():
            # 仅仅有一个key，还是Expert_ID说明这是为了执行删除操作
            delete_sql = "update texpert_info set is_delete=TRUE where expert_id=%s;"
            print(cursor.mogrify(delete_sql, args=data['Expert_ID']))
            delete_rows = cursor.execute(delete_sql, args=data['Expert_ID'])
            if delete_rows == 0:
                raise pymysql.MySQLError('数据删除失败，可能是SQL语句或数据类型的问题')
            else:
                conn.commit()
                return "SUCCESS"

        # 先查Expert_id对应的数据是否存在
        select_sql = "select expert_id from texpert_info where expert_id=%s;"
        select_rows = cursor.execute(select_sql, args=[data['Expert_ID']])
        if select_rows == 0:
            print('Expert_id不存在，创建新记录')
            insert_sql = "insert into texpert_info(id, expert_id, fieldId, name, college, education, degree, job, " \
                         "professional, subject, phonenumber, address, research, department, title, email, field, " \
                         "project) value(%(Expert_ID)s, %(Expert_ID)s, %(fieldID)s, %(name)s, %(college)s, " \
                         "%(education)s, %(degree)s, %(job)s, %(professional)s, %(subject)s, %(phone)s, %(address)s," \
                         "%(research)s, %(department)s, %(title)s, %(email)s, %(field)s, %(project)s);"
            # 这里字典的顺序和SQL语句参数的顺序一致，所以可以直接.values()
            # print(cursor.mogrify(insert_sql, args=data))
            rows = cursor.execute(insert_sql, args=data)
            if rows == 0:
                raise pymysql.MySQLError('数据插入失败，可能是SQL语句或数据类型的问题')
        else:
            print('Expert_id存在，更新记录')
            update_sql = "update texpert_info set fieldId=%(fieldID)s, name=%(name)s, college=%(college)s, " \
                         "education=%(education)s, degree=%(degree)s, job=%(job)s, professional=%(professional)s, " \
                         "subject=%(subject)s, phonenumber=%(phone)s, address=%(address)s, research=%(research)s, " \
                         "department=%(department)s, title=%(title)s, email=%(email)s, field=%(field)s, " \
                         "project=%(project)s, is_delete=FALSE where expert_id=%(Expert_ID)s;"
            # 查看生成的SQL语句
            # print(cursor.mogrify(update_sql, args=data))
            rows = cursor.execute(update_sql, args=data)
            print('修改影响的行数：', rows)
        conn.commit()
        return "SUCCESS"

    except Exception as e:
        print('数据库操作异常，开始回滚', e)
        conn.rollback()
        return "ERROR"
    finally:
        # 5：关闭游标、关闭连接
        cursor.close()
        conn.close()


# 修改数据库is_download字段为1
def modify_is_download(expert_id, is_delete):
    try:
        # 2.建立与数据库连接 TODO: 用户名root，密码x5
        conn = pymysql.connect(host="localhost",
                               user="root",
                               password="x5",
                               database="expertserverbig",
                               port=3306,
                               charset="UTF8")

        # 3.获取游标
        cursor = conn.cursor()
    except ConnectionError as e:
        print("数据库连接异常 或 获取游标异常！", e)
        return "数据保存失败"
    try:
        if is_delete:
            update_sql = "update texpert_info set is_download=TRUE where expert_id=%s;"
        else:
            update_sql = "update texpert_info set is_download=FALSE where expert_id=%s;"
        print(cursor.mogrify(update_sql, args=expert_id))
        update_rows = cursor.execute(update_sql, args=expert_id)
        conn.commit()
        return "SUCCESS"
    except Exception as e:
        print('数据库操作异常，开始回滚', e)
        conn.rollback()
        return "ERROR"
    finally:
        # 5：关闭游标、关闭连接
        cursor.close()
        conn.close()


# 解压压缩包
def unzip_file(zip_src, dst_dir):
    """
    1.解压zip文件
    2.如果是图片保存到data/pic目录下；
    3.如果是Excel，解析数据并将数据保存到数据库，然后文件放到data目录下
    :param zip_src: zip文件的全路径
    :param dst_dir: 要解压到的目的文件夹
    :param filename: 压缩包名字
    :return:
    """
    # 判断是否是压缩文件
    if not zipfile.is_zipfile(zip_src):
        return "请上传正确的zip压缩文件"
    fz = zipfile.ZipFile(zip_src, "r")
    fz.extractall(dst_dir)
    fz.close()

    # 查找解压后的文件，Excel就解析数据保存到数据库
    find_excel_file(dst_dir)
    return "SUCCESS"

# 解压rar类型压缩包(经测试，rar文件不可用)
def unrar_file(rar_src, dst_dir):
    # 判断是否是压缩文件
    if not rarfile.is_rarfile(rar_src):
        return "请上传正确的rar压缩文件"
    rar = rarfile.RarFile(rar_src, "r")
    rar.extractall(dst_dir)
    rar.close()

    # 查找解压后的文件，Excel就解析数据保存到数据库
    find_excel_file(dst_dir)
    return "SUCCESS"

# 解压gz类型压缩包
def ungz_file(rar_src, dst_dir):
    gz = gzip.GzipFile(rar_src, "r")
    open(dst_dir, "w+").write(gz.read())
    gz.close()

    # 查找解压后的文件，Excel就解析数据保存到数据库
    find_excel_file(dst_dir)
    return "SUCCESS"

# 解压tar类型压缩包
def untar_file(tar_src, dst_dir):
    # 判断是否是压缩文件
    if not tarfile.is_tarfile(tar_src):
        return "请上传正确的tar压缩文件"
    tar = tarfile.open(tar_src, "r")
    tar.extractall(dst_dir)
    tar.close()

    # 查找解压后的文件，Excel就解析数据保存到数据库
    find_excel_file(dst_dir)
    return "SUCCESS"

def find_excel_file(path):
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    files = os.listdir(path)
    for file in files:
        this_path = os.path.join(path, file)
        if os.path.isdir(this_path):
            find_excel_file(this_path)
        else:
            ret_list = file.rsplit(".", maxsplit=1)
            if ret_list[1] in ['xlsx', 'xls']:
                print(f"{file}是Excel文件，开始解析里面的数据。")
                # 开始读取里面的数据保存到数据库
                file_path = os.path.join(path, file)
                # 解析Excel模板中的数据
                doc = parse_excel(file_path)
                # 保存到数据库 优化 1.如果expert_id已经存在则更新数据 2.如果expert_id不存在则新增数据
                result1 = save_to_mysql(doc)
                result2 = modify_is_download(doc['Expert_ID'], True)
                if (result1 == 'ERROR') or (result2 == "REEOR"):
                    print("!!!!!!!上传文件，并保存到数据库或修改可下载字段失败！")
                    return "ERROR"
                else:
                    try:
                        # 先删除原来的
                        es.delete(index=my_index, doc_type=my_doc_type, id=doc['Expert_ID'])
                    except Exception as e:
                        print("!!!!!!!!ES中不存在对应ID的文档，此处为新增操作，接下来直接新增ES即可。", e)
                    # 再创建新的
                    res = es.index(index=my_index, doc_type=my_doc_type, body=doc, id=doc['Expert_ID'])
                    if (res['result'] != "created") and (res['result'] != "updated"):
                        print("!!!!!!!数据库新增或更新数据后，ES中更新失败！")
                # 重命名文件
                new_file_name = str(doc['Expert_ID']) + '.xlsx'
                new_file_path = os.path.join(BASE_DIR, 'data/excel/' + new_file_name)
                # 判断对应名字的文件是否存在，如果已存在，删除原来的文件保存新的文件
                if os.path.exists(new_file_path):
                    os.remove(new_file_path)

                os.renames(file_path, new_file_path)
            elif ret_list[1] in ['docx', 'doc']:
                print(f"{file}是Word文件，开始解析里面的数据。")
                # 开始读取里面的数据保存到数据库
                file_path = os.path.join(path, file)
                # 解析Word模板中的数据
                doc = parse_word(file_path)
                # 保存到数据库 优化 1.如果expert_id已经存在则更新数据 2.如果expert_id不存在则新增数据
                result1 = save_to_mysql(doc)
                result2 = modify_is_download(doc['Expert_ID'], True)
                if (result1 == 'ERROR') or (result2 == "REEOR"):
                    print("!!!!!!!上传World文件，并保存到数据库或修改可下载字段失败！")
                    return "ERROR"
                else:
                    try:
                        # 先删除原来的
                        es.delete(index=my_index, doc_type=my_doc_type, id=doc['Expert_ID'])
                    except Exception as e:
                        print("!!!!!!!!ES中不存在对应ID的文档，此处为新增操作，接下来直接新增ES即可。", e)
                    # 再创建新的
                    res = es.index(index=my_index, doc_type=my_doc_type, body=doc, id=doc['Expert_ID'])
                    if (res['result'] != "created") and (res['result'] != "updated"):
                        print("!!!!!!!数据库新增或更新数据后，ES中更新失败！")
                # 重命名文件
                new_file_name = str(doc['Expert_ID']) + '.xlsx'
                new_file_path = os.path.join(BASE_DIR, 'data/excel/' + new_file_name)
                # 判断对应名字的文件是否存在，如果已存在，删除原来的文件保存新的文件
                if os.path.exists(new_file_path):
                    os.remove(new_file_path)

                os.renames(file_path, new_file_path)
            else:
                # 假设其它都是图片，保存到pictures目录
                new_file_path = os.path.join(BASE_DIR, 'serverWeb/static/src/pic/' + file)
                # 判断对应名字的文件是否存在，如果已存在，删除原来的文件保存新的文件
                if os.path.exists(new_file_path):
                    os.remove(new_file_path)
                os.renames(this_path, new_file_path)


# if __name__ == '__main__':
#     mysql_data = get_data_from_mysql()
#     # TODO: FieldID的索引时long类型，不能是字符串
#     import_index_data(mysql_data)